#!/usr/bin/env python3
"""
GobLab Ficha Generator
Generates committee review sheets (fichas resumidas) from source URLs.

Usage:
    python3 generate_ficha.py URL1 URL2 ...
    python3 generate_ficha.py --file urls.txt
    python3 generate_ficha.py --pdf path/to/document.pdf
"""

import sys
import os
import json
import argparse
import textwrap
from datetime import datetime
from pathlib import Path

import requests
from bs4 import BeautifulSoup
import pdfplumber
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── CONFIG ────────────────────────────────────────────────────────────────────

MODEL = "claude-opus-4-8"
MAX_CONTENT_CHARS = 12000   # truncate fetched content before sending to Claude
OUTPUT_DIR = Path("fichas_output")

TEAL = RGBColor(0x00, 0x94, 0x99)     # GobLab teal
DARK = RGBColor(0x1A, 0x1A, 0x1A)


# ── CONTENT FETCHING ──────────────────────────────────────────────────────────

_FETCH_HEADERS = {
    "User-Agent":      "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "es-CL,es;q=0.9,en;q=0.8",
}


def fetch_url(url: str) -> str:
    """Fetch text content from a URL (HTML or PDF)."""
    try:
        r = requests.get(url, headers=_FETCH_HEADERS, timeout=20, allow_redirects=True)
        r.raise_for_status()
        content_type = r.headers.get("content-type", "")
        if "pdf" in content_type or url.lower().endswith(".pdf"):
            return extract_pdf_from_bytes(r.content)
        soup = BeautifulSoup(r.text, "html.parser")
        # Remove scripts, styles, nav
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return text
    except Exception as e:
        return f"[ERROR fetching {url}: {e}]"


def extract_pdf_from_path(path: str) -> str:
    """Extract text from a local PDF file."""
    text_parts = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
    except Exception as e:
        return f"[ERROR reading PDF {path}: {e}]"
    return "\n".join(text_parts)


def extract_pdf_from_bytes(data: bytes) -> str:
    """Extract text from PDF bytes."""
    import io
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
    except Exception as e:
        return f"[ERROR parsing PDF bytes: {e}]"
    return "\n".join(text_parts)


def truncate(text: str, max_chars: int = MAX_CONTENT_CHARS) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + f"\n\n[... contenido truncado a {max_chars} caracteres ...]"


# ── CLAUDE EXTRACTION ─────────────────────────────────────────────────────────

SYSTEM_PROMPT = """Eres un asistente especializado en el Repositorio de Algoritmos Públicos del GobLab UAI de la Universidad Adolfo Ibáñez (Chile). Tu tarea es analizar documentos sobre sistemas algorítmicos del sector público chileno y extraer información estructurada para una ficha resumida que será presentada al Comité Editorial.

Debes responder ÚNICAMENTE con un objeto JSON válido con exactamente estas claves. Si no hay información suficiente para un campo, usa "Sin información disponible".

{
  "nombre": "Nombre formal del sistema o nombre tentativo descriptivo",
  "objetivo_sistema": "Breve descripción del problema que busca resolver (1-2 oraciones)",
  "decision_automatizada": "Descripción de cuál es el proceso de decisión que se automatiza o semiautomatiza",
  "tipo_automatizacion": "Automático | Semiautomático | Sin información disponible",
  "usa_datos_personales": "Sí | No | Sin información disponible",
  "institucion_publica": "Nombre completo de la institución pública asociada",
  "unidad": "Nombre de la subdirección, departamento o unidad dentro de la institución (si se menciona)",
  "ejecutor": "Nombre de la organización o empresa que elaboró o desarrolla el sistema",
  "fuentes_apa": ["Fuente 1 en formato APA", "Fuente 2 en formato APA"],
  "notas_revisor": "Observaciones relevantes para el Comité: ambigüedades, información faltante importante, o señales de que el sistema podría no cumplir criterios de publicación"
}"""


def extract_ficha_fields(content_blocks: list[dict], client: anthropic.Anthropic) -> dict:
    """
    content_blocks: [{"url": ..., "text": ...}, ...]
    Returns parsed ficha dict.
    """
    sources_text = ""
    for i, block in enumerate(content_blocks, 1):
        label = block.get("url") or block.get("path") or f"Fuente {i}"
        sources_text += f"\n\n--- FUENTE {i}: {label} ---\n{truncate(block['text'])}"

    user_message = f"""Analiza el siguiente contenido y extrae la información para la ficha resumida del Repositorio de Algoritmos Públicos.

El contenido proviene de {len(content_blocks)} fuente(s):
{sources_text}

Responde únicamente con el JSON estructurado solicitado."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )

    raw = response.content[0].text.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Retry once with a higher token budget and an explicit length warning
        try:
            retry_resp = client.messages.create(
                model=MODEL,
                max_tokens=6000,
                system=SYSTEM_PROMPT,
                messages=[
                    {"role": "user", "content": user_message},
                    {"role": "assistant", "content": raw},
                    {"role": "user", "content":
                     "Tu respuesta anterior fue cortada. Por favor, repite el JSON completo y válido "
                     "desde el principio, sin omitir ningún campo."},
                ],
            )
            raw2 = retry_resp.content[0].text.strip()
            if raw2.startswith("```"):
                raw2 = raw2.split("```")[1]
                if raw2.startswith("json"):
                    raw2 = raw2[4:]
            return json.loads(raw2)
        except Exception:
            return {"_parse_error": raw}


# ── WORD DOCUMENT GENERATION ──────────────────────────────────────────────────

def add_heading_row(table, text, colspan=2):
    """Add a teal header row spanning the table."""
    row = table.add_row()
    cell = row.cells[0]
    cell.merge(row.cells[1])
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "009499")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_field_row(table, label, value):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_cell.text = label
    label_run = label_cell.paragraphs[0].runs[0]
    label_run.bold = True
    label_run.font.size = Pt(9)

    value_cell.text = str(value) if value else "Sin información disponible"
    for run in value_cell.paragraphs[0].runs:
        run.font.size = Pt(9)

    # Light background on label column
    tc = label_cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "E8F7F7")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def build_docx(fichas: list[dict], source_map: list[dict], output_path: Path):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Fichas Resumidas para el Comité Editorial", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = TEAL

    subtitle = doc.add_paragraph("Repositorio de Algoritmos Públicos — GobLab UAI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_p = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    for i, (ficha, sources) in enumerate(zip(fichas, source_map), 1):
        if i > 1:
            doc.add_page_break()

        # Algorithm number heading
        algo_title = doc.add_heading(
            f"Algoritmo #{i}: {ficha.get('nombre', 'Sin nombre')}", level=2
        )
        for run in algo_title.runs:
            run.font.color.rgb = TEAL

        # Sources used
        src_labels = [s.get("url") or s.get("path") or "Fuente" for s in sources]
        src_p = doc.add_paragraph(f"Fuentes analizadas: {' | '.join(src_labels)}")
        for run in src_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            run.italic = True

        doc.add_paragraph()

        # Check for parse error
        if "_parse_error" in ficha:
            doc.add_paragraph(
                "⚠ Error al parsear respuesta de Claude. Respuesta cruda:",
                style="Intense Quote",
            )
            doc.add_paragraph(ficha["_parse_error"])
            continue

        # Main fields table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        # Column widths
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)

        add_heading_row(table, "IDENTIFICACIÓN DEL SISTEMA")
        add_field_row(table, "Nombre", ficha.get("nombre"))
        add_field_row(table, "Objetivo del Sistema", ficha.get("objetivo_sistema"))
        add_field_row(table, "Decisión automatizada", ficha.get("decision_automatizada"))
        add_field_row(table, "Automático / Semiautomático", ficha.get("tipo_automatizacion"))
        add_field_row(table, "Utiliza datos personales", ficha.get("usa_datos_personales"))

        add_heading_row(table, "INSTITUCIÓN Y EJECUTOR")
        add_field_row(table, "Institución Pública", ficha.get("institucion_publica"))
        add_field_row(table, "Dirección / Unidad", ficha.get("unidad"))
        add_field_row(table, "Ejecutor", ficha.get("ejecutor"))

        # Sources
        add_heading_row(table, "FUENTES")
        sources_apa = ficha.get("fuentes_apa", [])
        if sources_apa:
            sources_text = "\n".join(f"[{j+1}] {s}" for j, s in enumerate(sources_apa))
        else:
            sources_text = "Sin información disponible"
        add_field_row(table, "Links / Fuentes (APA)", sources_text)

        # Reviewer notes
        if ficha.get("notas_revisor"):
            add_heading_row(table, "NOTAS PARA EL COMITÉ")
            add_field_row(table, "Observaciones", ficha.get("notas_revisor"))

        # Set column widths (must be done after rows are added)
        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(11)

        doc.add_paragraph()

        # Approval checkbox row
        status_p = doc.add_paragraph()
        status_p.add_run("Decisión del Comité:  ").bold = True
        status_p.add_run("☐ Aprobado    ☐ Monitoreo    ☐ Rechazado    ☐ Pendiente más info")
        status_p.paragraph_format.space_before = Pt(6)

    doc.save(output_path)
    print(f"\n✓ Fichas guardadas en: {output_path}")


# ── MAIN ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Genera fichas resumidas para el Comité Editorial del Repositorio de Algoritmos Públicos."
    )
    parser.add_argument("urls", nargs="*", help="URLs a analizar")
    parser.add_argument("--file", "-f", help="Archivo de texto con una URL por línea")
    parser.add_argument("--pdf", "-p", nargs="+", help="Archivos PDF locales a analizar")
    parser.add_argument("--output", "-o", default=None, help="Nombre del archivo Word de salida")
    parser.add_argument(
        "--group",
        action="store_true",
        help="Agrupa todas las fuentes en UNA sola ficha (útil cuando múltiples URLs son sobre el mismo algoritmo)",
    )
    args = parser.parse_args()

    # Collect all sources
    all_sources = []  # list of {"url"/"path": ..., "text": ...}

    for url in args.urls:
        print(f"  Fetching {url} ...")
        text = fetch_url(url)
        all_sources.append({"url": url, "text": text})

    if args.file:
        with open(args.file) as f:
            for line in f:
                url = line.strip()
                if url and not url.startswith("#"):
                    print(f"  Fetching {url} ...")
                    text = fetch_url(url)
                    all_sources.append({"url": url, "text": text})

    for pdf_path in (args.pdf or []):
        print(f"  Reading PDF {pdf_path} ...")
        text = extract_pdf_from_path(pdf_path)
        all_sources.append({"path": pdf_path, "text": text})

    if not all_sources:
        parser.print_help()
        sys.exit(1)

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        print("ERROR: Set ANTHROPIC_API_KEY environment variable.")
        sys.exit(1)

    client = anthropic.Anthropic(api_key=api_key)

    # Decide grouping: --group bundles everything into one ficha
    # Default: each source = one ficha (unless --group)
    if args.group:
        groups = [all_sources]          # one group, all sources
    else:
        groups = [[s] for s in all_sources]   # one group per source

    OUTPUT_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = args.output or f"fichas_{timestamp}.docx"
    output_path = OUTPUT_DIR / out_name

    fichas = []
    for j, group in enumerate(groups, 1):
        label = group[0].get("url") or group[0].get("path") or f"grupo {j}"
        print(f"\nAnalizando con Claude: {label} ...")
        ficha = extract_ficha_fields(group, client)
        fichas.append(ficha)

        # Print summary to console
        if "_parse_error" not in ficha:
            print(f"  → {ficha.get('nombre', '?')}")
            print(f"    Institución: {ficha.get('institucion_publica', '?')}")
            print(f"    Tipo: {ficha.get('tipo_automatizacion', '?')}  |  Datos personales: {ficha.get('usa_datos_personales', '?')}")

    build_docx(fichas, groups, output_path)


if __name__ == "__main__":
    main()
